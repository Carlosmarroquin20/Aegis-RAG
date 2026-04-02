"""
DOCX parser using python-docx.

Returns one RawDocument per logical section (heading-delimited).
Falls back to a single RawDocument containing all paragraphs if no headings
are present (flat documents without structure).

Heading detection uses python-docx's built-in style name matching ("Heading 1",
"Heading 2", etc.) rather than font size heuristics, which are unreliable.
"""

from __future__ import annotations

import structlog

from aegis.domain.models.ingestion import RawDocument, SupportedMimeType
from aegis.domain.ports.document_parser import DocumentParserPort, ParseError

logger = structlog.get_logger(__name__)


class DocxParser(DocumentParserPort):
    def supported_mime_types(self) -> frozenset[str]:
        return frozenset({SupportedMimeType.APPLICATION_DOCX})

    def parse(self, content: bytes, filename: str) -> list[RawDocument]:
        try:
            import io

            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ParseError(filename, "python-docx is not installed") from exc

        try:
            doc = DocxDocument(io.BytesIO(content))
        except Exception as exc:  # noqa: BLE001
            raise ParseError(filename, f"could not open DOCX: {exc}") from exc

        sections = self._extract_sections(doc, filename)
        if not sections:
            raise ParseError(filename, "no readable text found in document")

        return sections

    @staticmethod
    def _extract_sections(doc: object, filename: str) -> list[RawDocument]:
        """
        Splits document into sections at Heading 1/2 boundaries.
        Each section becomes a RawDocument with the heading as metadata.
        """
        from docx import Document as DocxDocument  # type: ignore[attr-defined]

        doc_typed: DocxDocument = doc  # type: ignore[assignment]
        raw_docs: list[RawDocument] = []
        current_heading: str | None = None
        current_lines: list[str] = []

        def flush(heading: str | None, lines: list[str]) -> None:
            text = "\n".join(lines).strip()
            if text:
                raw_docs.append(
                    RawDocument(
                        content=text,
                        source=filename,
                        section=heading,
                        metadata={
                            "format": "docx",
                            **({"section": heading} if heading else {}),
                        },
                    )
                )

        for para in doc_typed.paragraphs:
            style_name: str = para.style.name if para.style else ""
            text = para.text.strip()

            if style_name.startswith("Heading"):
                flush(current_heading, current_lines)
                current_heading = text or style_name
                current_lines = []
            elif text:
                current_lines.append(text)

        flush(current_heading, current_lines)
        return raw_docs
